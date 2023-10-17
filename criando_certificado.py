# Programa que lê uma seleção de lutadores e gera um certificado que compareceu na sua luta.
# Formatação a estudar.
from docx import Document
from docx.shared import Cm

document = Document()

document.add_heading('Certificado de luta', 0)

document.add_paragraph('O medo de perder tira a vontade de ganhar.', style='Intense Quote')

document.add_paragraph(
f'Certificado dado ao lutador @nome, dos pesos-@categoria, por completar sua luta no evento.',
)

document.add_picture('assets/pngwing.com.png', width=Cm(5))

document.add_page_break()

document.save(f'assets/certificado_base2.docx')